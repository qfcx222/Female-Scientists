# -*- coding: utf-8 -*-
"""
女科学家传记情感倾向与学术关系可视化分析 - 中文专业版
针对中文文本的完整解决方案
"""
import jieba
import jieba.posseg as pseg
import jieba.analyse
import pandas as pd
import numpy as np
import re
import os
import warnings
import json
import pickle
import hashlib
from datetime import datetime
from typing import Dict, List, Tuple, Optional, Any, Set
import logging
from tqdm import tqdm
from collections import defaultdict, Counter
from concurrent.futures import ThreadPoolExecutor, as_completed
import gc
from docx import Document
from sklearn.preprocessing import MinMaxScaler, StandardScaler
from sklearn.feature_extraction.text import TfidfVectorizer, CountVectorizer
from sklearn.decomposition import LatentDirichletAllocation, NMF
from sklearn.metrics import silhouette_score, calinski_harabasz_score
from sklearn.cluster import KMeans
from sklearn.manifold import TSNE
import matplotlib.pyplot as plt
import seaborn as sns
import networkx as nx
import matplotlib
from matplotlib import font_manager
import community as community_louvain
from scipy import sparse
from scipy.stats import zscore

# 设置matplotlib中文字体
try:
    # 尝试多种中文字体路径
    font_paths = [
        "C:\\Windows\\Fonts\\simhei.ttf",  # Windows黑体
        "C:\\Windows\\Fonts\\msyh.ttc",  # Windows微软雅黑
        "C:\\Windows\\Fonts\\simsun.ttc",  # Windows宋体
        "/System/Library/Fonts/PingFang.ttc",  # Mac苹方
        "/usr/share/fonts/truetype/droid/DroidSansFallbackFull.ttf",  # Linux
    ]

    for font_path in font_paths:
        if os.path.exists(font_path):
            font_prop = font_manager.FontProperties(fname=font_path)
            matplotlib.rcParams['font.sans-serif'] = [font_prop.get_name()]
            matplotlib.rcParams['axes.unicode_minus'] = False
            break
    else:
        # 如果找不到字体文件，尝试使用系统已安装的字体
        matplotlib.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'SimSun']
        matplotlib.rcParams['axes.unicode_minus'] = False
except Exception as e:
    print(f"字体设置警告: {e}")

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('chinese_analysis.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

warnings.filterwarnings('ignore')


# 初始化jieba，添加专业词汇
def initialize_jieba():
    """初始化jieba分词器，添加专业词汇"""
    # 添加科学家人名（可根据需要扩展）
    scientists = [
        '屠呦呦', '张弥曼', '颜宁', '庄小威', '李飞飞',
        '王小云', '刘若川', '吴健雄', '林巧稚', '何泽慧'
    ]
    for name in scientists:
        jieba.add_word(name, freq=1000, tag='nr')

    # 添加学术词汇
    academic_words = [
        '生物学家', '化学家', '物理学家', '数学家', '计算机科学家',
        '研究员', '教授', '博士生导师', '院士', '学术委员会',
        '国家重点实验室', '国家自然科学基金', '学术论文', '学术会议',
        '研究成果', '科研项目', '学术交流', '国际合作', '学术报告'
    ]
    for word in academic_words:
        jieba.add_word(word, freq=500, tag='n')

    # 加载自定义词典（如果有）
    custom_dict_path = "custom_dict.txt"
    if os.path.exists(custom_dict_path):
        jieba.load_userdict(custom_dict_path)
        logger.info(f"已加载自定义词典: {custom_dict_path}")


initialize_jieba()

# 中文停用词扩展（学术专用）
CHINESE_STOPWORDS = set([
    # 基础停用词
    '的', '了', '在', '是', '我', '有', '和', '就', '不', '人', '都', '一', '一个', '上', '也', '很', '到', '说', '要',
    '去', '你',
    '会', '着', '没有', '看', '好', '自己', '这', '那', '他', '她', '它', '我们', '你们', '他们', '她们', '它们',
    '而', '但', '且', '或', '并', '及', '与', '于', '对', '对于', '关于', '至于', '以及', '及其',

    # 学术停用词
    '研究', '工作', '论文', '发表', '实验', '数据', '结果', '方法', '分析', '理论',
    '模型', '系统', '技术', '过程', '领域', '科学', '学术', '学者', '科研',
    '大学', '学院', '研究所', '实验室', '中心', '机构', '部门', '课题组',
    '项目', '基金', '经费', '资助', '奖励', '奖项', '荣誉', '称号', '学位',
    '博士', '硕士', '学士', '博士后', '教授', '副教授', '讲师', '研究员',
    '团队', '合作', '协作', '交流', '会议', '报告', '演讲', '讲座',
    '发现', '发明', '创新', '创造', '贡献', '影响', '意义', '价值',
    '方面', '部分', '内容', '问题', '挑战', '困难', '局限', '未来',
    '目前', '现在', '过去', '当前', '近年来', '近年来来', '近年来来',
])

# 中文情感词典（扩展版）
CHINESE_SENTIMENT_DICT = {
    'positive': {
        '优秀': 0.3, '杰出': 0.35, '卓越': 0.4, '突出': 0.25, '显著': 0.3,
        '重要': 0.2, '关键': 0.25, '核心': 0.25, '主要': 0.15, '重大': 0.35,
        '创新': 0.3, '创造': 0.25, '发明': 0.35, '突破': 0.4, '开创': 0.35,
        '成功': 0.25, '成就': 0.3, '成果': 0.2, '贡献': 0.25, '影响': 0.2,
        '著名': 0.3, '知名': 0.25, '声望': 0.2, '声誉': 0.2, '荣誉': 0.25,
        '积极': 0.2, '乐观': 0.15, '热情': 0.15, '专注': 0.15, '执着': 0.2,
        '严谨': 0.2, '认真': 0.15, '细致': 0.15, '耐心': 0.1, '坚持': 0.2,
        '聪明': 0.15, '智慧': 0.2, '才华': 0.2, '天赋': 0.25, '才能': 0.15,
        '领导': 0.2, '指导': 0.15, '培养': 0.15, '教育': 0.1, '启发': 0.15,
        '合作': 0.1, '团队': 0.1, '交流': 0.05, '国际': 0.1, '全球': 0.1,
    },
    'negative': {
        '困难': 0.2, '挑战': 0.15, '问题': 0.1, '障碍': 0.25, '阻碍': 0.2,
        '失败': 0.3, '挫折': 0.25, '失望': 0.2, '遗憾': 0.15, '不足': 0.1,
        '局限': 0.15, '限制': 0.1, '缺陷': 0.25, '缺点': 0.2, '弱点': 0.2,
        '压力': 0.15, '紧张': 0.1, '焦虑': 0.2, '担忧': 0.15, '恐惧': 0.25,
        '批评': 0.2, '质疑': 0.15, '怀疑': 0.1, '反对': 0.2, '争议': 0.25,
        '偏见': 0.3, '歧视': 0.35, '不公平': 0.3, '不平等': 0.25, '障碍': 0.2,
        '孤独': 0.2, '孤立': 0.25, '排斥': 0.3, '边缘': 0.25, '忽视': 0.2,
        '复杂': 0.1, '混乱': 0.15, '矛盾': 0.2, '冲突': 0.25, '竞争': 0.15,
        '牺牲': 0.25, '放弃': 0.2, '离开': 0.15, '失去': 0.2, '病痛': 0.3,
        '死亡': 0.4, '疾病': 0.3, '健康': -0.1, '治疗': -0.05, '恢复': 0.1,
    }
}

# 学术主题关键词（中文）
CHINESE_ACADEMIC_TOPICS = {
    'research_methodology': ['方法', '方法论', '技术', '手段', '途径', '实验', '设计'],
    'theoretical_framework': ['理论', '框架', '模型', '概念', '原理', '假设', '范式'],
    'empirical_study': ['数据', '样本', '证据', '实证', '观察', '测量', '统计'],
    'career_development': ['职业', '生涯', '职位', '晋升', '成就', '荣誉', '奖励'],
    'collaboration_network': ['合作', '团队', '网络', '伙伴', '交流', '国际', '跨学科'],
    'scientific_discovery': ['发现', '发明', '创新', '突破', '成果', '贡献', '影响'],
    'academic_recognition': ['奖项', '荣誉', '称号', '院士', '杰出', '优秀', '表彰'],
    'education_training': ['教育', '培养', '指导', '教学', '学习', '训练', '导师'],
    'institutional_context': ['机构', '大学', '学院', '研究所', '实验室', '中心', '部门'],
    'scientific_impact': ['影响', '贡献', '意义', '价值', '重要性', '里程碑', '开创性']
}

# 中文关系提取模式
CHINESE_RELATIONSHIP_PATTERNS = {
    'collaborators': [
        (re.compile(r'(?:与|和|同|跟)([^，。；！？、]{2,8})(?:合作|协作|共同)'), 3),
        (re.compile(r'(?:合作|协作)(?:者|伙伴|方)(?:[：:])?([^，。；！？、]{2,8})'), 3),
        (re.compile(r'(?:与|和|同|跟)([^，。；！？、]{2,8})(?:一起|一同|联合)'), 2),
        (re.compile(r'(?:联合|共同)(?:研究|发表|完成)(?:[：:])?([^，。；！？、]{2,8})'), 3),
    ],
    'advisors': [
        (re.compile(r'(?:导师|指导教师|博士生导师)(?:[：:])?([^，。；！？、]{2,8})'), 5),
        (re.compile(r'(?:师从|师承|受教于)([^，。；！？、]{2,8})'), 4),
        (re.compile(r'(?:在)([^，。；！？、]{2,8})(?:教授|导师)(?:指导下)'), 4),
        (re.compile(r'(?:博士|硕士|研究生)(?:阶段)?(?:导师)(?:[：:])?([^，。；！？、]{2,8})'), 5),
    ],
    'students': [
        (re.compile(r'(?:学生|弟子|研究生)(?:[：:])?([^，。；！？、]{2,8})'), 4),
        (re.compile(r'(?:指导|培养)(?:了)?([^，。；！？、]{2,8})'), 3),
        (re.compile(r'(?:博士生|硕士生|研究生)(?:[：:])?([^，。；！？、]{2,8})'), 4),
        (re.compile(r'(?:作为)([^，。；！？、]{2,8})(?:的导师)'), 4),
    ],
    'colleagues': [
        (re.compile(r'(?:同事|同僚|同仁)(?:[：:])?([^，。；！？、]{2,8})'), 3),
        (re.compile(r'(?:与|和|同|跟)([^，。；！？、]{2,8})(?:共事|同事)'), 2),
        (re.compile(r'(?:工作伙伴|团队成员)(?:[：:])?([^，。；！？、]{2,8})'), 3),
    ],
    'institutions': [
        (re.compile(r'(?:大学|学院|研究所|实验室|中心)(?:[：:])?([^，。；！？、]{4,15})'), 2),
        (re.compile(r'(?:任职于|工作于|就读于)([^，。；！？、]{4,15})'), 3),
        (re.compile(r'(?:在)([^，。；！？、]{4,15})(?:工作|学习|研究)'), 2),
        (re.compile(r'(?:毕业于)([^，。；！？、]{4,15})'), 2),
    ]
}


class ChineseScientistBiographyAnalyzer:
    """中文科学家传记分析器"""

    def __init__(self, input_path: str, output_folder: str = "chinese_results"):
        """
        初始化中文分析器

        Args:
            input_path: Word文档文件夹路径或文件路径
            output_folder: 输出文件夹路径
        """
        self.input_path = input_path
        self.output_folder = output_folder

        # 初始化数据结构
        self.biographies = {}
        self.df = None
        self.sentiment_df = None
        self.relationship_graph = None
        self.relationship_data = None
        self.topics = None
        self.topic_distributions = None

        # 初始化工具
        self.scaler = MinMaxScaler()
        self.std_scaler = StandardScaler()

        # 质量评估
        self.quality_metrics = {}
        self.performance_stats = {}

        # 缓存
        self._cache = {}

        # 确保输出目录存在
        os.makedirs(self.output_folder, exist_ok=True)

        logger.info(f"初始化中文分析器: input={input_path}, output={output_folder}")

    def _cache_key(self, func_name: str, *args) -> str:
        """生成缓存键"""
        arg_str = "_".join(str(arg) for arg in args)
        key = f"{func_name}_{hashlib.md5(arg_str.encode()).hexdigest()[:8]}"
        return key

    def _get_cached(self, key: str):
        """获取缓存结果"""
        return self._cache.get(key)

    def _set_cached(self, key: str, value):
        """设置缓存"""
        self._cache[key] = value

    def extract_text_from_docx(self, docx_path: str) -> str:
        """从Word文档提取文本（支持中文）"""
        cache_key = self._cache_key("extract_text", docx_path)
        cached = self._get_cached(cache_key)
        if cached is not None:
            return cached

        try:
            doc = Document(docx_path)
            text_parts = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())

            text = "\n".join(text_parts)

            # 缓存结果
            self._set_cached(cache_key, text)
            return text

        except Exception as e:
            logger.error(f"提取Word文档失败 {docx_path}: {e}")
            return ""

    def clean_chinese_text(self, text: str) -> str:
        """中文文本清洗"""
        if not text:
            return ""

        # 去除特殊字符和多余空白
        text = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9\s，。；！？、：%（）【】《》"\'\-]', '', text)
        text = re.sub(r'\s+', ' ', text)

        # 处理标点符号
        text = re.sub(r'([，。；！？、])\1+', r'\1', text)

        # 去除网址和邮箱
        text = re.sub(r'http\S+|www\S+|https\S+', '', text)
        text = re.sub(r'\S+@\S+', '', text)

        # 去除数字和字母（可选，根据需求调整）
        # text = re.sub(r'[a-zA-Z0-9]+', '', text)

        return text.strip()

    def segment_chinese_text(self, text: str, use_pos: bool = False) -> Tuple[List[str], List[str]]:
        """
        中文分词和短语提取

        Args:
            text: 待分词文本
            use_pos: 是否使用词性标注

        Returns:
            tokens: 分词结果
            phrases: 关键短语
        """
        cache_key = self._cache_key("segment_text", text, use_pos)
        cached = self._get_cached(cache_key)
        if cached is not None:
            return cached

        if not text:
            result = ([], [])
            self._set_cached(cache_key, result)
            return result

        # 分词
        if use_pos:
            # 使用词性标注
            words = pseg.cut(text)
            tokens = []
            for word, flag in words:
                if word not in CHINESE_STOPWORDS and len(word) > 1:
                    # 保留名词、动词、形容词
                    if flag.startswith(('n', 'v', 'a')):
                        tokens.append(word)
        else:
            # 普通分词
            tokens = [word for word in jieba.cut(text)
                      if word not in CHINESE_STOPWORDS and len(word) > 1]

        # 提取关键短语（TF-IDF）
        phrases = []
        if tokens:
            # 使用jieba的TF-IDF关键词提取
            try:
                keywords = jieba.analyse.extract_tags(
                    ' '.join(tokens),
                    topK=20,
                    withWeight=False,
                    allowPOS=('n', 'v', 'a')
                )
                phrases = list(keywords)
            except:
                # 回退方法：提取2-3gram
                n = len(tokens)
                for i in range(n - 1):
                    phrase = tokens[i] + tokens[i + 1]
                    if len(phrase) >= 4:
                        phrases.append(phrase)

                if n >= 3:
                    for i in range(n - 2):
                        phrase = tokens[i] + tokens[i + 1] + tokens[i + 2]
                        if len(phrase) >= 6:
                            phrases.append(phrase)

        result = (tokens, phrases)
        self._set_cached(cache_key, result)
        return result

    def split_chinese_sentences(self, text: str) -> List[str]:
        """中文分句"""
        if not text:
            return []

        # 使用中文标点分句
        sentences = re.split(r'[。！？；\n]+', text)
        sentences = [s.strip() for s in sentences if len(s.strip()) > 5]

        # 进一步分割长句子（以逗号分割）
        final_sentences = []
        for sentence in sentences:
            if len(sentence) > 50:
                # 如果句子太长，按逗号分割
                sub_sentences = re.split(r'[，,]+', sentence)
                sub_sentences = [s.strip() for s in sub_sentences if len(s.strip()) > 10]
                final_sentences.extend(sub_sentences)
            else:
                final_sentences.append(sentence)

        return final_sentences

    def analyze_chinese_sentiment(self, text: str) -> Dict[str, float]:
        """
        中文情感分析（基于词典的方法）

        Returns:
            包含情感分数和置信度的字典
        """
        if not text:
            return {'score': 0, 'confidence': 0, 'positive_words': [], 'negative_words': []}

        # 分词
        words = jieba.cut(text)
        words = [w for w in words if w not in CHINESE_STOPWORDS and len(w) > 1]

        # 统计情感词
        positive_words = []
        negative_words = []
        positive_score = 0
        negative_score = 0

        for word in words:
            if word in CHINESE_SENTIMENT_DICT['positive']:
                positive_words.append(word)
                positive_score += CHINESE_SENTIMENT_DICT['positive'][word]
            elif word in CHINESE_SENTIMENT_DICT['negative']:
                negative_words.append(word)
                negative_score += CHINESE_SENTIMENT_DICT['negative'][word]

        # 计算基础分数（-1到1）
        total_score = positive_score - negative_score

        # 归一化
        word_count = len(words)
        if word_count > 0:
            normalized_score = total_score / (word_count * 0.5)  # 调整系数
            normalized_score = max(-1, min(1, normalized_score))
        else:
            normalized_score = 0

        # 计算置信度
        sentiment_words = len(positive_words) + len(negative_words)
        if word_count > 0:
            confidence = sentiment_words / word_count
        else:
            confidence = 0

        # 调整置信度
        confidence = max(0.1, min(1.0, confidence * 2))

        return {
            'score': normalized_score,
            'confidence': confidence,
            'positive_words': positive_words[:10],
            'negative_words': negative_words[:10],
            'sentiment_words_count': sentiment_words,
            'total_words_count': word_count
        }

    def analyze_chinese_sentiment_advanced(self, text: str) -> Dict[str, float]:
        """
        高级中文情感分析（结合句法和上下文）
        """
        sentences = self.split_chinese_sentences(text)

        if not sentences:
            return {'score': 0, 'confidence': 0, 'sentence_scores': []}

        sentence_scores = []
        sentence_confidences = []

        for sentence in sentences:
            # 基础情感分析
            sentiment_result = self.analyze_chinese_sentiment(sentence)

            # 考虑否定词
            if any(neg_word in sentence for neg_word in ['不', '没', '无', '未', '非']):
                sentiment_result['score'] *= -0.5  # 否定词反转

            # 考虑程度副词
            degree_words = {
                '非常': 1.5, '十分': 1.5, '极其': 2.0, '极度': 2.0,
                '很': 1.3, '较': 1.2, '稍微': 0.8, '略微': 0.8,
                '特别': 1.5, '尤其': 1.4, '相当': 1.3
            }

            for word, factor in degree_words.items():
                if word in sentence:
                    sentiment_result['score'] *= factor
                    sentiment_result['confidence'] *= 1.1
                    break

            # 限制分数范围
            sentiment_result['score'] = max(-1, min(1, sentiment_result['score']))

            sentence_scores.append(sentiment_result['score'])
            sentence_confidences.append(sentiment_result['confidence'])

        # 计算整体分数（加权平均）
        if sentence_scores:
            avg_score = np.mean(sentence_scores)
            avg_confidence = np.mean(sentence_confidences)
        else:
            avg_score = 0
            avg_confidence = 0

        return {
            'score': avg_score,
            'confidence': avg_confidence,
            'sentence_scores': sentence_scores,
            'sentence_count': len(sentences)
        }

    def extract_chinese_relationships(self, text: str, scientist_name: str) -> Dict:
        """
        提取中文文本中的学术关系

        Args:
            text: 文本内容
            scientist_name: 科学家姓名

        Returns:
            关系字典
        """
        relationships = defaultdict(lambda: defaultdict(int))

        # 策略1：模式匹配
        self._extract_relationships_by_patterns(text, scientist_name, relationships)

        # 策略2：命名实体识别（简化版）
        self._extract_relationships_by_ner(text, scientist_name, relationships)

        # 策略3：依存句法分析（简化版）
        self._extract_relationships_by_syntax(text, scientist_name, relationships)

        # 合并和过滤结果
        return self._merge_and_filter_chinese_relationships(relationships, scientist_name)

    def _extract_relationships_by_patterns(self, text: str, scientist_name: str, relationships: Dict):
        """基于模式提取关系"""
        for rel_type, patterns in CHINESE_RELATIONSHIP_PATTERNS.items():
            for pattern, weight in patterns:
                matches = pattern.finditer(text)
                for match in matches:
                    if match.groups():
                        entity = match.group(1).strip()
                        if (entity and
                                entity != scientist_name and
                                2 <= len(entity) <= 15):
                            relationships[rel_type][entity] += weight

    def _extract_relationships_by_ner(self, text: str, scientist_name: str, relationships: Dict):
        """基于命名实体识别提取关系（简化版）"""
        # 使用jieba的词性标注来识别人名和机构名
        words = pseg.cut(text)

        persons = []
        organizations = []

        for word, flag in words:
            if flag == 'nr' and word not in CHINESE_STOPWORDS and len(word) >= 2:
                persons.append(word)
            elif flag == 'nt' or any(org_word in word for org_word in ['大学', '学院', '研究所', '实验室']):
                organizations.append(word)

        # 在上下文中寻找关系线索
        sentences = self.split_chinese_sentences(text)

        for sentence in sentences:
            sentence_lower = sentence

            # 合作者关系
            if any(keyword in sentence for keyword in ['合作', '协作', '共同']):
                for person in persons:
                    if person in sentence and person != scientist_name:
                        relationships['collaborators'][person] += 2

            # 导师关系
            if any(keyword in sentence for keyword in ['导师', '师从', '受教']):
                for person in persons:
                    if person in sentence and person != scientist_name:
                        relationships['advisors'][person] += 3

            # 学生关系
            if any(keyword in sentence for keyword in ['学生', '指导', '培养']):
                for person in persons:
                    if person in sentence and person != scientist_name:
                        relationships['students'][person] += 2

            # 机构关系
            for org in organizations:
                if org in sentence:
                    relationships['institutions'][org] += 1

    def _extract_relationships_by_syntax(self, text: str, scientist_name: str, relationships: Dict):
        """基于句法分析提取关系（简化版）"""
        # 这里可以集成更复杂的句法分析工具，如LTP、HanLP等
        # 目前使用简单的规则匹配

        sentences = self.split_chinese_sentences(text)

        for sentence in sentences:
            # 简单的依存关系规则
            if '与' in sentence or '和' in sentence or '同' in sentence:
                # 尝试提取并列关系
                parts = re.split(r'[与和同]', sentence)
                if len(parts) >= 2:
                    for part in parts:
                        # 提取可能的人名（2-4个字符）
                        possible_names = re.findall(r'[\u4e00-\u9fa5]{2,4}', part)
                        for name in possible_names:
                            if name != scientist_name and len(name) >= 2:
                                relationships['colleagues'][name] += 1

    def _merge_and_filter_chinese_relationships(self, relationships: Dict, scientist_name: str) -> Dict:
        """合并和过滤中文关系"""
        filtered = {}

        thresholds = {
            'collaborators': 1,
            'advisors': 1,
            'students': 1,
            'colleagues': 1,
            'institutions': 1
        }

        for rel_type, rel_dict in relationships.items():
            threshold = thresholds.get(rel_type, 1)

            # 合并变体（简单的去重）
            merged = defaultdict(int)
            for name, weight in rel_dict.items():
                if weight >= threshold:
                    # 标准化名称
                    normalized_name = self._normalize_chinese_name(name)
                    if normalized_name:
                        merged[normalized_name] = max(merged[normalized_name], weight)

            # 排序并限制数量
            sorted_items = sorted(merged.items(), key=lambda x: x[1], reverse=True)[:10]
            filtered[rel_type] = sorted_items

        return filtered

    def _normalize_chinese_name(self, name: str) -> str:
        """标准化中文名称"""
        if not name:
            return ""

        # 去除常见前缀和后缀
        name = name.strip()
        prefixes = ['教授', '博士', '老师', '先生', '女士', '研究员', '院士']
        suffixes = ['教授', '博士', '老师', '先生', '女士', '研究员', '院士']

        for prefix in prefixes:
            if name.startswith(prefix):
                name = name[len(prefix):].strip()

        for suffix in suffixes:
            if name.endswith(suffix):
                name = name[:-len(suffix)].strip()

        return name

    def load_and_preprocess_biographies(self) -> bool:
        """加载并预处理所有中文传记"""
        logger.info("开始加载和预处理中文Word文档...")
        start_time = datetime.now()

        # 获取文档文件
        docx_files = self._get_docx_files()

        if not docx_files:
            logger.error("未找到Word文档")
            return False

        # 处理每个文档
        scientist_data = []

        for docx_file in tqdm(docx_files, desc="处理文档"):
            result = self._process_single_chinese_biography(docx_file)
            if result:
                scientist_data.append(result)

        if not scientist_data:
            logger.error("没有成功提取任何有效文本")
            return False

        # 构建DataFrame
        processed_data = []

        for data in scientist_data:
            # 分词和短语提取
            tokens, phrases = self.segment_chinese_text(data['cleaned_text'], use_pos=True)

            # 计算文本统计指标
            word_count = len(tokens)
            sentence_count = len(data['sentences'])
            unique_words = len(set(tokens))

            # 计算文本复杂度
            ttr = unique_words / word_count if word_count > 0 else 0
            avg_sent_len = np.mean([len(s) for s in data['sentences']]) if data['sentences'] else 0

            processed_data.append({
                'scientist': data['name'],
                'full_text': data['cleaned_text'],
                'raw_text': data['raw_text'],
                'sentences': data['sentences'],
                'tokens': tokens,
                'phrases': phrases,
                'word_count': word_count,
                'sentence_count': sentence_count,
                'unique_word_count': unique_words,
                'avg_word_length': np.mean([len(t) for t in tokens]) if tokens else 0,
                'ttr': ttr,
                'avg_sentence_length': avg_sent_len,
                'sentence_length_std': np.std([len(s) for s in data['sentences']]) if len(data['sentences']) > 1 else 0,
                'file_path': data['file_path']
            })

        self.df = pd.DataFrame(processed_data)

        # 记录质量指标
        self.quality_metrics['data_loading'] = {
            'total_files': len(docx_files),
            'valid_files': len(scientist_data),
            'success_rate': len(scientist_data) / len(docx_files),
            'avg_text_length': self.df['word_count'].mean(),
            'avg_sentence_count': self.df['sentence_count'].mean(),
            'avg_ttr': self.df['ttr'].mean()
        }

        # 性能统计
        self.performance_stats['load_time'] = (datetime.now() - start_time).total_seconds()

        logger.info(
            f"预处理完成: {len(self.df)}/{len(docx_files)} 篇传记，耗时: {self.performance_stats['load_time']:.1f}秒")
        return True

    def _get_docx_files(self) -> List[str]:
        """获取Word文档文件列表"""
        docx_files = []

        if os.path.isdir(self.input_path):
            for root, dirs, files in os.walk(self.input_path):
                for f in files:
                    if f.lower().endswith(('.docx', '.doc')) and not f.startswith('~$'):
                        docx_files.append(os.path.join(root, f))
        elif os.path.isfile(self.input_path) and self.input_path.lower().endswith(('.docx', '.doc')):
            docx_files = [self.input_path]

        return docx_files

    def _process_single_chinese_biography(self, docx_file: str) -> Optional[Dict]:
        """处理单篇中文传记"""
        scientist_name = os.path.splitext(os.path.basename(docx_file))[0]
        text = self.extract_text_from_docx(docx_file)

        if not text or len(text.strip()) < 100:
            logger.warning(f"文本过短或为空: {scientist_name}")
            return None

        # 清洗文本
        cleaned_text = self.clean_chinese_text(text)

        # 分句
        sentences = self.split_chinese_sentences(cleaned_text)

        # 基础验证
        if len(sentences) < 3:
            logger.warning(f"句子数量不足: {scientist_name}")
            return None

        return {
            'name': scientist_name,
            'raw_text': text,
            'cleaned_text': cleaned_text,
            'sentences': sentences,
            'file_path': docx_file
        }

    def analyze_sentiment_for_all(self) -> bool:
        """为所有科学家进行情感分析"""
        if self.df is None:
            logger.error("请先加载传记数据")
            return False

        logger.info("开始中文情感分析...")
        start_time = datetime.now()

        sentiment_results = []

        for idx, row in tqdm(self.df.iterrows(), total=len(self.df), desc="情感分析"):
            scientist = row['scientist']
            text = row['full_text']
            sentences = row['sentences']

            if not text:
                sentiment_results.append(self._create_default_sentiment_result(scientist))
                continue

            # 进行高级情感分析
            sentiment_result = self.analyze_chinese_sentiment_advanced(text)

            # 句子级分析
            sentence_analyses = []
            positive_sentences = []
            negative_sentences = []
            neutral_sentences = []

            for sentence in sentences[:100]:  # 限制句子数量
                if len(sentence.strip()) < 5:
                    continue

                # 句子情感分析
                sent_sentiment = self.analyze_chinese_sentiment(sentence)
                score = sent_sentiment['score']

                # 分类
                if score > 0.1:
                    sentiment_label = 'positive'
                    positive_sentences.append(sentence)
                elif score < -0.1:
                    sentiment_label = 'negative'
                    negative_sentences.append(sentence)
                else:
                    sentiment_label = 'neutral'
                    neutral_sentences.append(sentence)

                sentence_analyses.append({
                    'sentence': sentence,
                    'score': score,
                    'sentiment_label': sentiment_label,
                    'confidence': sent_sentiment['confidence']
                })

            # 计算统计指标
            total_sent = len(sentence_analyses)
            if total_sent > 0:
                pos_ratio = len(positive_sentences) / total_sent
                neg_ratio = len(negative_sentences) / total_sent
                neu_ratio = len(neutral_sentences) / total_sent

                # 情感多样性（熵）
                sentiment_counts = Counter([s['sentiment_label'] for s in sentence_analyses])
                entropy = -sum((count / total_sent) * np.log2(count / total_sent)
                               for count in sentiment_counts.values() if count > 0)
            else:
                pos_ratio = neg_ratio = neu_ratio = entropy = 0

            sentiment_results.append({
                'scientist': scientist,
                'avg_sentiment': sentiment_result['score'],
                'sentiment_confidence': sentiment_result['confidence'],
                'sentiment_entropy': entropy,
                'positive_ratio': pos_ratio,
                'negative_ratio': neg_ratio,
                'neutral_ratio': neu_ratio,
                'positive_count': len(positive_sentences),
                'negative_count': len(negative_sentences),
                'neutral_count': len(neutral_sentences),
                'total_sentences': total_sent,
                'sentence_analyses': sentence_analyses[:50],  # 限制数量
                'positive_words': sentiment_result.get('positive_words', []),
                'negative_words': sentiment_result.get('negative_words', [])
            })

        self.sentiment_df = pd.DataFrame(sentiment_results)

        # 记录质量指标
        self.quality_metrics['sentiment_analysis'] = {
            'avg_confidence': self.sentiment_df['sentiment_confidence'].mean(),
            'avg_sentiment': self.sentiment_df['avg_sentiment'].mean(),
            'sentiment_diversity': self.sentiment_df['sentiment_entropy'].mean(),
            'data_points': self.sentiment_df['total_sentences'].sum()
        }

        # 性能统计
        self.performance_stats['sentiment_time'] = (datetime.now() - start_time).total_seconds()

        logger.info(f"情感分析完成，耗时: {self.performance_stats['sentiment_time']:.1f}秒")
        logger.info(f"平均情感分数: {self.sentiment_df['avg_sentiment'].mean():.3f}")

        return True

    def _create_default_sentiment_result(self, scientist: str) -> Dict:
        """创建默认情感分析结果"""
        return {
            'scientist': scientist,
            'avg_sentiment': 0,
            'sentiment_confidence': 0.5,
            'sentiment_entropy': 0,
            'positive_ratio': 0,
            'negative_ratio': 0,
            'neutral_ratio': 1,
            'positive_count': 0,
            'negative_count': 0,
            'neutral_count': 0,
            'total_sentences': 0,
            'sentence_analyses': [],
            'positive_words': [],
            'negative_words': []
        }

    def build_relationship_network(self) -> bool:
        """构建中文学术关系网络"""
        if self.df is None:
            logger.error("请先加载传记数据")
            return False

        logger.info("开始构建中文关系网络...")
        start_time = datetime.now()

        # 创建图
        G = nx.Graph()
        relationship_data = []

        for idx, row in tqdm(self.df.iterrows(), total=len(self.df), desc="构建网络"):
            scientist = row['scientist']

            # 添加科学家节点
            sentiment_score = 0
            if self.sentiment_df is not None:
                sentiment_row = self.sentiment_df[self.sentiment_df['scientist'] == scientist]
                if not sentiment_row.empty:
                    sentiment_score = sentiment_row.iloc[0]['avg_sentiment']

            G.add_node(scientist,
                       type='scientist',
                       sentiment=sentiment_score,
                       word_count=row['word_count'],
                       sentence_count=row['sentence_count'])

            # 提取关系
            relationships = self.extract_chinese_relationships(
                row['full_text'], scientist
            )

            # 记录关系数据
            rel_summary = {
                'scientist': scientist,
                'total_relationships': 0,
                'relationships': {}
            }

            # 添加关系边
            for rel_type, rel_items in relationships.items():
                rel_summary['relationships'][rel_type] = rel_items
                rel_summary['total_relationships'] += len(rel_items)

                for entity, weight in rel_items:
                    # 添加实体节点
                    if entity not in G.nodes():
                        G.add_node(entity, type=rel_type[:-1] if rel_type.endswith('s') else rel_type,
                                   weight=weight)

                    # 添加边（带权重）
                    if G.has_edge(scientist, entity):
                        # 更新权重
                        current_weight = G[scientist][entity].get('weight', 0)
                        G[scientist][entity]['weight'] = max(current_weight, weight)
                    else:
                        G.add_edge(scientist, entity,
                                   relationship=rel_type,
                                   weight=weight,
                                   confidence=min(1.0, weight / 5))

            relationship_data.append(rel_summary)

        self.relationship_graph = G
        self.relationship_data = relationship_data

        # 计算网络指标
        if G.number_of_nodes() > 0:
            density = nx.density(G)
            avg_degree = sum(dict(G.degree()).values()) / G.number_of_nodes()

            # 记录质量指标
            self.quality_metrics['relationship_network'] = {
                'nodes': G.number_of_nodes(),
                'edges': G.number_of_edges(),
                'density': density,
                'avg_degree': avg_degree,
                'scientists_with_relations': sum(1 for d in relationship_data if d['total_relationships'] > 0),
                'avg_relations_per_scientist': np.mean([d['total_relationships'] for d in relationship_data])
            }

        # 性能统计
        self.performance_stats['network_time'] = (datetime.now() - start_time).total_seconds()

        logger.info(f"关系网络构建完成，耗时: {self.performance_stats['network_time']:.1f}秒")
        logger.info(f"网络统计: {G.number_of_nodes()}节点, {G.number_of_edges()}边, 密度: {density:.3f}")

        return True

    def perform_chinese_topic_modeling(self, num_topics: int = 5, method: str = 'lda') -> bool:
        """中文主题建模"""
        if self.df is None or len(self.df) < 3:
            logger.warning("数据不足，跳过主题建模")
            return False

        logger.info(f"开始中文主题建模 ({method.upper()})...")
        start_time = datetime.now()

        # 准备文档
        documents = []
        valid_indices = []

        for idx, row in self.df.iterrows():
            if row['tokens'] and len(row['tokens']) >= 10:
                # 使用分词结果构建文档
                doc_text = ' '.join(row['tokens'][:500])  # 限制长度
                documents.append(doc_text)
                valid_indices.append(idx)

        if len(documents) < 3:
            logger.warning("有效文档不足，跳过主题建模")
            return False

        # 向量化（中文使用CountVectorizer或TfidfVectorizer）
        vectorizer = CountVectorizer(
            max_df=0.9,
            min_df=2,
            max_features=1000,
            token_pattern=r'(?u)\b\w+\b'  # 匹配中文字符
        )

        try:
            X = vectorizer.fit_transform(documents)

            # 选择算法
            if method == 'lda':
                model = LatentDirichletAllocation(
                    n_components=num_topics,
                    random_state=42,
                    max_iter=20,
                    learning_method='online'
                )
            elif method == 'nmf':
                model = NMF(
                    n_components=num_topics,
                    random_state=42,
                    max_iter=200
                )
            else:
                logger.error(f"不支持的主题建模方法: {method}")
                return False

            # 训练模型
            topic_distributions = model.fit_transform(X)

            # 提取主题
            feature_names = vectorizer.get_feature_names_out()
            topics = []

            for topic_idx in range(num_topics):
                if method == 'lda':
                    top_indices = model.components_[topic_idx].argsort()[-15:][::-1]
                else:
                    top_indices = model.components_[topic_idx].argsort()[-15:][::-1]

                top_words = [feature_names[i] for i in top_indices]
                top_weights = [model.components_[topic_idx][i] for i in top_indices]

                # 识别主题类型
                topic_type = self._identify_chinese_topic_type(top_words)

                topics.append({
                    'topic_id': topic_idx,
                    'top_words': top_words,
                    'word_weights': top_weights,
                    'topic_type': topic_type,
                    'topic_name': f"主题{topic_idx}: {topic_type}",
                    'coherence_score': self._calculate_chinese_topic_coherence(top_words, documents)
                })

            # 分配主题给科学家
            self._assign_topics_to_scientists_chinese(topic_distributions, topics, valid_indices)

            self.topics = topics
            self.topic_distributions = topic_distributions

            # 记录质量指标
            avg_coherence = np.mean([t['coherence_score'] for t in topics])
            self.quality_metrics['topic_modeling'] = {
                'method': method,
                'num_topics': num_topics,
                'avg_coherence': avg_coherence,
                'avg_topic_confidence': self.df[
                    'topic_confidence'].mean() if 'topic_confidence' in self.df.columns else 0
            }

            # 性能统计
            self.performance_stats['topic_time'] = (datetime.now() - start_time).total_seconds()

            logger.info(f"主题建模完成，耗时: {self.performance_stats['topic_time']:.1f}秒")
            logger.info(f"主题质量: 平均一致性 {avg_coherence:.3f}")

            return True

        except Exception as e:
            logger.error(f"主题建模失败: {e}")
            return False

    def _identify_chinese_topic_type(self, top_words: List[str]) -> str:
        """识别中文主题类型"""
        top_words_str = "".join(top_words)

        for topic_name, keywords in CHINESE_ACADEMIC_TOPICS.items():
            keyword_count = sum(1 for keyword in keywords if keyword in top_words_str)
            if keyword_count >= 2:
                return topic_name

        # 基于关键词推断
        research_words = ['研究', '实验', '数据', '分析', '方法']
        career_words = ['职业', '成就', '荣誉', '奖项', '晋升']
        collab_words = ['合作', '团队', '交流', '国际', '网络']

        if any(word in top_words_str for word in research_words):
            return '研究方法'
        elif any(word in top_words_str for word in career_words):
            return '职业发展'
        elif any(word in top_words_str for word in collab_words):
            return '合作网络'
        else:
            return '一般学术'

    def _calculate_chinese_topic_coherence(self, top_words: List[str], documents: List[str]) -> float:
        """计算中文主题一致性"""
        if len(top_words) < 2 or len(documents) < 2:
            return 0.0

        # 计算词对共现
        co_occurrence = 0
        total_pairs = 0

        # 只计算前5个词的相关性
        top_n = min(5, len(top_words))

        for i in range(top_n):
            for j in range(i + 1, top_n):
                word1, word2 = top_words[i], top_words[j]
                pair_count = sum(1 for doc in documents if word1 in doc and word2 in doc)
                co_occurrence += pair_count
                total_pairs += 1

        if total_pairs > 0:
            max_possible = len(documents) * total_pairs
            return co_occurrence / max_possible if max_possible > 0 else 0
        else:
            return 0.0

    def _assign_topics_to_scientists_chinese(self, topic_distributions, topics, valid_indices):
        """分配主题给科学家（中文）"""
        # 初始化列
        self.df['dominant_topic'] = -1
        self.df['dominant_topic_name'] = ''
        self.df['topic_distribution'] = None
        self.df['topic_confidence'] = 0.0

        # 分配主题
        for idx, matrix_idx in enumerate(valid_indices):
            if idx < len(topic_distributions):
                dist = topic_distributions[idx]
                if len(dist) > 0:
                    dominant_idx = dist.argmax()
                    confidence = dist[dominant_idx]

                    self.df.at[matrix_idx, 'dominant_topic'] = dominant_idx
                    self.df.at[matrix_idx, 'dominant_topic_name'] = topics[dominant_idx]['topic_type']
                    self.df.at[matrix_idx, 'topic_distribution'] = dist.tolist()
                    self.df.at[matrix_idx, 'topic_confidence'] = confidence

    def create_visualizations(self):
        """创建中文可视化图表"""
        logger.info("创建可视化图表...")

        try:
            # 1. 情感分析可视化
            self._plot_sentiment_analysis_chinese()

            # 2. 关系网络可视化
            self._plot_relationship_network_chinese()

            # 3. 主题建模可视化
            if self.topics is not None:
                self._plot_topic_modeling_chinese()

            # 4. 综合可视化
            self._plot_comprehensive_chinese()

            logger.info("可视化图表创建完成")

        except Exception as e:
            logger.error(f"可视化创建失败: {e}")

    def _plot_sentiment_analysis_chinese(self):
        """绘制中文情感分析结果"""
        if self.sentiment_df is None:
            return

        fig, axes = plt.subplots(2, 2, figsize=(15, 12))

        # 1. 情感分数分布
        ax1 = axes[0, 0]
        scientists = self.sentiment_df['scientist']
        scores = self.sentiment_df['avg_sentiment']

        colors = ['#4CAF50' if s > 0.1 else '#F44336' if s < -0.1 else '#2196F3' for s in scores]
        bars = ax1.bar(range(len(scientists)), scores, color=colors, alpha=0.7)

        ax1.axhline(y=0, color='black', linestyle='-', alpha=0.3)
        ax1.set_xticks(range(len(scientists)))
        ax1.set_xticklabels(scientists, rotation=45, ha='right')
        ax1.set_ylabel('情感分数')
        ax1.set_title('科学家情感分数分布', fontsize=12, fontweight='bold')

        # 2. 情感成分比例
        ax2 = axes[0, 1]
        pos_ratios = self.sentiment_df['positive_ratio']
        neg_ratios = self.sentiment_df['negative_ratio']
        neu_ratios = self.sentiment_df['neutral_ratio']

        x = range(len(scientists))
        ax2.bar(x, pos_ratios, label='积极', color='#4CAF50', alpha=0.8)
        ax2.bar(x, neg_ratios, bottom=pos_ratios, label='消极', color='#F44336', alpha=0.8)
        ax2.bar(x, neu_ratios, bottom=pos_ratios + neg_ratios, label='中性', color='#2196F3', alpha=0.8)

        ax2.set_xticks(x)
        ax2.set_xticklabels(scientists, rotation=45, ha='right')
        ax2.set_ylabel('比例')
        ax2.set_title('情感成分比例', fontsize=12, fontweight='bold')
        ax2.legend(bbox_to_anchor=(1.05, 1), loc='upper left')

        # 3. 情感置信度
        ax3 = axes[1, 0]
        confidences = self.sentiment_df['sentiment_confidence']

        ax3.hist(confidences, bins=15, color='#9C27B0', alpha=0.7, edgecolor='black')
        ax3.axvline(confidences.mean(), color='red', linestyle='--',
                    label=f'平均值: {confidences.mean():.3f}')

        ax3.set_xlabel('置信度')
        ax3.set_ylabel('频次')
        ax3.set_title('情感分析置信度分布', fontsize=12, fontweight='bold')
        ax3.legend()
        ax3.grid(True, alpha=0.3)

        # 4. 情感多样性
        ax4 = axes[1, 1]
        entropy = self.sentiment_df['sentiment_entropy']

        bars = ax4.bar(range(len(entropy)), entropy, color='#FF9800', alpha=0.7)

        ax4.set_xticks(range(len(entropy)))
        ax4.set_xticklabels(scientists, rotation=45, ha='right')
        ax4.set_ylabel('熵值')
        ax4.set_title('情感多样性（熵）', fontsize=12, fontweight='bold')

        # 添加数值标签
        for bar, val in zip(bars, entropy):
            ax4.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.01,
                     f'{val:.2f}', ha='center', va='bottom', fontsize=8)

        plt.suptitle('中文科学家传记情感分析结果', fontsize=14, fontweight='bold')
        plt.tight_layout()

        # 保存图表
        save_path = os.path.join(self.output_folder, "sentiment_analysis_chinese.png")
        plt.savefig(save_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.show()

        logger.info(f"情感分析图表已保存至: {save_path}")

    def _plot_relationship_network_chinese(self):
        """绘制中文关系网络"""
        if self.relationship_graph is None:
            return

        plt.figure(figsize=(14, 12))

        G = self.relationship_graph

        # 节点分类
        scientist_nodes = [n for n, d in G.nodes(data=True) if d.get('type') == 'scientist']
        other_nodes = [n for n in G.nodes() if n not in scientist_nodes]

        # 布局算法
        pos = nx.spring_layout(G, k=2, iterations=100, seed=42)

        # 绘制节点
        nx.draw_networkx_nodes(G, pos, nodelist=scientist_nodes,
                               node_color='#FF6B6B', node_size=500,
                               alpha=0.9, label='科学家')

        if other_nodes:
            nx.draw_networkx_nodes(G, pos, nodelist=other_nodes,
                                   node_color='#45B7D1', node_size=300,
                                   alpha=0.7, label='其他实体')

        # 绘制边（按关系类型着色）
        edge_colors = []
        edge_widths = []

        for u, v, data in G.edges(data=True):
            rel_type = data.get('relationship', 'unknown')
            weight = data.get('weight', 1)

            if 'collaborat' in rel_type:
                edge_colors.append('#FF4500')  # 合作 - 橙红
            elif 'advisor' in rel_type:
                edge_colors.append('#2E8B57')  # 导师 - 深绿
            elif 'student' in rel_type:
                edge_colors.append('#98FB98')  # 学生 - 浅绿
            elif 'colleague' in rel_type:
                edge_colors.append('#1E90FF')  # 同事 - 蓝色
            elif 'institution' in rel_type:
                edge_colors.append('#9370DB')  # 机构 - 紫色
            else:
                edge_colors.append('#A9A9A9')  # 其他 - 灰色

            edge_widths.append(max(0.5, weight * 0.5))

        nx.draw_networkx_edges(G, pos, edge_color=edge_colors,
                               width=edge_widths, alpha=0.6)

        # 绘制标签（只显示科学家和重要节点）
        labels = {}
        for node in G.nodes():
            if node in scientist_nodes or G.degree(node) >= 3:
                labels[node] = node if len(node) <= 10 else node[:8] + '...'

        nx.draw_networkx_labels(G, pos, labels, font_size=9)

        plt.title('中文科学家学术关系网络', fontsize=14, fontweight='bold')
        plt.legend(loc='upper left')
        plt.axis('off')

        # 添加网络统计信息
        stats_text = f"""
        网络统计:
        节点总数: {G.number_of_nodes()}
        边总数: {G.number_of_edges()}
        网络密度: {nx.density(G):.4f}
        平均节点度: {sum(dict(G.degree()).values()) / G.number_of_nodes():.2f}
        科学家节点: {len(scientist_nodes)}
        """

        plt.figtext(0.02, 0.02, stats_text, fontsize=10,
                    bbox=dict(boxstyle='round', facecolor='wheat', alpha=0.8))

        # 保存图表
        save_path = os.path.join(self.output_folder, "relationship_network_chinese.png")
        plt.savefig(save_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.show()

        logger.info(f"关系网络图表已保存至: {save_path}")

    def _plot_topic_modeling_chinese(self):
        """绘制中文主题建模结果"""
        if self.topics is None:
            return

        fig, axes = plt.subplots(1, 2, figsize=(15, 6))

        # 1. 主题关键词
        ax1 = axes[0]
        topic_data = []

        for topic in self.topics:
            topic_id = topic['topic_id']
            top_words = topic['top_words'][:10]  # 只显示前10个词
            words_text = ' '.join(top_words)
            topic_data.append((topic_id, topic['topic_type'], words_text))

        # 创建表格
        table_data = [[d[0], d[1], d[2]] for d in topic_data]
        columns = ['主题ID', '主题类型', '关键词']

        ax1.axis('tight')
        ax1.axis('off')
        table = ax1.table(cellText=table_data, colLabels=columns,
                          cellLoc='left', loc='center')
        table.auto_set_font_size(False)
        table.set_fontsize(9)
        table.scale(1, 1.5)

        ax1.set_title('主题建模结果', fontsize=12, fontweight='bold')

        # 2. 主题分布
        ax2 = axes[1]
        if 'dominant_topic' in self.df.columns:
            topic_counts = self.df['dominant_topic'].value_counts().sort_index()

            if len(topic_counts) > 0:
                colors_topic = plt.cm.Set3(np.linspace(0, 1, len(topic_counts)))
                bars = ax2.bar(range(len(topic_counts)), topic_counts.values, color=colors_topic)

                # 添加标签
                for bar, count in zip(bars, topic_counts.values):
                    ax2.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                             str(count), ha='center', va='bottom', fontsize=10)

                ax2.set_xticks(range(len(topic_counts)))
                ax2.set_xticklabels([f'主题{i}' for i in topic_counts.index], rotation=45)
                ax2.set_ylabel('科学家数量')
                ax2.set_title('主题分布', fontsize=12, fontweight='bold')
                ax2.grid(True, alpha=0.3, axis='y')

        plt.suptitle('中文科学家传记主题建模分析', fontsize=14, fontweight='bold')
        plt.tight_layout()

        # 保存图表
        save_path = os.path.join(self.output_folder, "topic_modeling_chinese.png")
        plt.savefig(save_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.show()

        logger.info(f"主题建模图表已保存至: {save_path}")

    def _plot_comprehensive_chinese(self):
        """绘制综合可视化图表"""
        fig = plt.figure(figsize=(18, 12))

        # 1. 文本长度分布
        if self.df is not None:
            ax1 = plt.subplot(2, 3, 1)
            word_counts = self.df['word_count']

            ax1.hist(word_counts, bins=15, color='#4CAF50', alpha=0.7, edgecolor='black')
            ax1.axvline(word_counts.mean(), color='red', linestyle='--',
                        label=f'平均: {word_counts.mean():.0f}词')

            ax1.set_xlabel('词数')
            ax1.set_ylabel('频次')
            ax1.set_title('传记文本长度分布', fontsize=11, fontweight='bold')
            ax1.legend()
            ax1.grid(True, alpha=0.3)

        # 2. 情感与文本长度关系
        if self.df is not None and self.sentiment_df is not None:
            ax2 = plt.subplot(2, 3, 2)

            merged_df = pd.merge(self.sentiment_df[['scientist', 'avg_sentiment']],
                                 self.df[['scientist', 'word_count']], on='scientist')

            ax2.scatter(merged_df['word_count'], merged_df['avg_sentiment'],
                        alpha=0.6, s=100, color='#FF6B6B')

            # 添加科学家标签
            for i, row in merged_df.iterrows():
                if i < 10:  # 只标注前10个
                    ax2.annotate(row['scientist'][:4],
                                 (row['word_count'], row['avg_sentiment']),
                                 fontsize=8, alpha=0.7)

            ax2.set_xlabel('文本长度（词数）')
            ax2.set_ylabel('情感分数')
            ax2.set_title('文本长度 vs 情感分数', fontsize=11, fontweight='bold')
            ax2.grid(True, alpha=0.3)

        # 3. 关系网络度分布
        if self.relationship_graph is not None:
            ax3 = plt.subplot(2, 3, 3)
            degrees = [d for n, d in self.relationship_graph.degree()]

            if degrees:
                ax3.hist(degrees, bins=min(15, len(set(degrees))),
                         color='#2196F3', alpha=0.7, edgecolor='black')
                ax3.axvline(np.mean(degrees), color='red', linestyle='--',
                            label=f'平均度: {np.mean(degrees):.2f}')

                ax3.set_xlabel('节点度')
                ax3.set_ylabel('频次')
                ax3.set_title('关系网络度分布', fontsize=11, fontweight='bold')
                ax3.legend()
                ax3.grid(True, alpha=0.3)

        # 4. 情感词云（模拟）
        ax4 = plt.subplot(2, 3, 4)

        # 收集所有情感词
        all_positive_words = []
        all_negative_words = []

        if self.sentiment_df is not None:
            for _, row in self.sentiment_df.iterrows():
                all_positive_words.extend(row.get('positive_words', []))
                all_negative_words.extend(row.get('negative_words', []))

        # 统计词频
        pos_counter = Counter(all_positive_words)
        neg_counter = Counter(all_negative_words)

        # 显示前10个积极词和消极词
        pos_words = pos_counter.most_common(10)
        neg_words = neg_counter.most_common(10)

        if pos_words or neg_words:
            # 创建简单条形图
            pos_labels = [w[0] for w in pos_words]
            pos_counts = [w[1] for w in pos_words]

            neg_labels = [w[0] for w in neg_words]
            neg_counts = [w[1] for w in neg_words]

            x_pos = range(len(pos_labels))
            ax4.barh(x_pos, pos_counts, color='#4CAF50', alpha=0.7, label='积极词')

            x_neg = range(len(neg_labels))
            ax4.barh([x + len(pos_labels) + 1 for x in x_neg], neg_counts,
                     color='#F44336', alpha=0.7, label='消极词')

            ax4.set_yticks(list(x_pos) + [x + len(pos_labels) + 1 for x in x_neg])
            ax4.set_yticklabels(pos_labels + neg_labels)
            ax4.set_xlabel('出现频次')
            ax4.set_title('情感词汇分布', fontsize=11, fontweight='bold')
            ax4.legend()

        # 5. 性能统计
        ax5 = plt.subplot(2, 3, 5)

        times = [v for k, v in self.performance_stats.items() if v > 0]
        labels = [k.replace('_time', '') for k, v in self.performance_stats.items() if v > 0]

        if times:
            colors_time = ['#FF6B6B', '#45B7D1', '#96CEB4', '#FFEAA7'][:len(times)]
            bars = ax5.bar(range(len(times)), times, color=colors_time, alpha=0.7)

            ax5.set_xticks(range(len(times)))
            ax5.set_xticklabels(labels, rotation=45)
            ax5.set_ylabel('时间（秒）')
            ax5.set_title('分析性能统计', fontsize=11, fontweight='bold')

            # 添加时间标签
            for bar, time in zip(bars, times):
                ax5.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                         f'{time:.1f}s', ha='center', va='bottom', fontsize=9)

        # 6. 质量评分
        ax6 = plt.subplot(2, 3, 6)

        if self.quality_metrics:
            dimensions = ['数据质量', '情感分析', '关系网络', '主题建模']
            scores = []

            if 'data_loading' in self.quality_metrics:
                scores.append(self.quality_metrics['data_loading'].get('success_rate', 0) * 10)
            else:
                scores.append(0)

            if 'sentiment_analysis' in self.quality_metrics:
                scores.append(self.quality_metrics['sentiment_analysis'].get('avg_confidence', 0) * 10)
            else:
                scores.append(0)

            if 'relationship_network' in self.quality_metrics:
                scores.append(min(10, self.quality_metrics['relationship_network'].get('avg_degree', 0) * 2))
            else:
                scores.append(0)

            if 'topic_modeling' in self.quality_metrics:
                scores.append(self.quality_metrics['topic_modeling'].get('avg_coherence', 0) * 10)
            else:
                scores.append(0)

            colors_quality = ['#4CAF50', '#2196F3', '#FF9800', '#9C27B0']
            bars = ax6.bar(range(len(scores)), scores, color=colors_quality, alpha=0.7)

            ax6.set_xticks(range(len(scores)))
            ax6.set_xticklabels(dimensions, rotation=45)
            ax6.set_ylabel('分数（0-10）')
            ax6.set_title('分析质量评估', fontsize=11, fontweight='bold')
            ax6.set_ylim(0, 10)

            # 添加分数标签
            for bar, score in zip(bars, scores):
                ax6.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.1,
                         f'{score:.1f}', ha='center', va='bottom', fontsize=9)

        plt.suptitle('中文科学家传记分析综合报告', fontsize=16, fontweight='bold', y=1.02)
        plt.tight_layout()

        # 保存图表
        save_path = os.path.join(self.output_folder, "comprehensive_analysis_chinese.png")
        plt.savefig(save_path, dpi=300, bbox_inches='tight', facecolor='white')
        plt.show()

        logger.info(f"综合图表已保存至: {save_path}")

    def export_results(self):
        """导出中文分析结果"""
        logger.info("导出中文分析结果...")

        # 1. Excel格式
        if self.df is not None and self.sentiment_df is not None:
            with pd.ExcelWriter(os.path.join(self.output_folder, "chinese_analysis_results.xlsx"),
                                engine='openpyxl') as writer:

                # 基础数据
                df_export = self.df.copy()
                if 'sentences' in df_export.columns:
                    df_export['sentences'] = df_export['sentences'].apply(
                        lambda x: ' | '.join(x) if isinstance(x, list) else x)
                if 'tokens' in df_export.columns:
                    df_export['tokens'] = df_export['tokens'].apply(lambda x: ' '.join(x) if isinstance(x, list) else x)

                df_export.to_excel(writer, sheet_name='基础数据', index=False)

                # 情感分析
                sentiment_export = self.sentiment_df.copy()
                if 'sentence_analyses' in sentiment_export.columns:
                    sentiment_export = sentiment_export.drop(columns=['sentence_analyses'])
                sentiment_export.to_excel(writer, sheet_name='情感分析', index=False)

                # 关系网络摘要
                if hasattr(self, 'relationship_data') and self.relationship_data:
                    rel_summary = []
                    for data in self.relationship_data:
                        rel_summary.append({
                            '科学家': data['scientist'],
                            '总关系数': data['total_relationships'],
                            '合作者': len(data['relationships'].get('collaborators', [])),
                            '导师': len(data['relationships'].get('advisors', [])),
                            '学生': len(data['relationships'].get('students', [])),
                            '同事': len(data['relationships'].get('colleagues', [])),
                            '机构': len(data['relationships'].get('institutions', []))
                        })
                    pd.DataFrame(rel_summary).to_excel(writer, sheet_name='关系网络', index=False)

                # 主题建模
                if self.topics is not None:
                    topics_df = pd.DataFrame([
                        {
                            '主题ID': t['topic_id'],
                            '主题类型': t['topic_type'],
                            '关键词': ', '.join(t['top_words'][:10]),
                            '一致性分数': t['coherence_score']
                        }
                        for t in self.topics
                    ])
                    topics_df.to_excel(writer, sheet_name='主题建模', index=False)

            logger.info("Excel格式结果已导出")

        # 2. JSON格式
        full_data = {
            'metadata': {
                'timestamp': datetime.now().isoformat(),
                'num_scientists': len(self.df) if self.df is not None else 0,
                'analysis_version': 'chinese_v1.0'
            },
            'quality_metrics': self.quality_metrics,
            'performance_stats': self.performance_stats,
            'summary': {
                'avg_sentiment': self.sentiment_df['avg_sentiment'].mean() if self.sentiment_df is not None else 0,
                'total_relations': self.relationship_graph.number_of_edges() if self.relationship_graph is not None else 0,
                'num_topics': len(self.topics) if self.topics is not None else 0
            }
        }

        with open(os.path.join(self.output_folder, "chinese_analysis_summary.json"), 'w',
                  encoding='utf-8') as f:
            json.dump(full_data, f, indent=2, ensure_ascii=False)

        # 3. 网络数据
        if self.relationship_graph is not None:
            try:
                nx.write_gexf(self.relationship_graph,
                              os.path.join(self.output_folder, "chinese_academic_network.gexf"))
                logger.info("网络数据（GEXF格式）已导出")
            except Exception as e:
                logger.warning(f"GEXF导出失败: {e}")

        # 4. 文本摘要
        self._export_chinese_summary()

        logger.info(f"所有结果已导出至: {self.output_folder}")

    def _export_chinese_summary(self):
        """导出中文文本摘要"""
        summary = []
        summary.append("=" * 60)
        summary.append("中文女科学家传记分析摘要")
        summary.append("=" * 60)
        summary.append("")

        if self.df is not None:
            summary.append(f"分析科学家数量: {len(self.df)}")
            summary.append(f"总词数: {self.df['word_count'].sum():,}")
            summary.append(f"总句数: {self.df['sentence_count'].sum():,}")
            summary.append(f"平均文本长度: {self.df['word_count'].mean():.0f} 词")
            summary.append("")

        if self.sentiment_df is not None:
            avg_sentiment = self.sentiment_df['avg_sentiment'].mean()
            if avg_sentiment > 0.1:
                sentiment_label = "积极"
            elif avg_sentiment < -0.1:
                sentiment_label = "消极"
            else:
                sentiment_label = "中性"

            summary.append(f"整体情感倾向: {sentiment_label} (分数: {avg_sentiment:.3f})")
            summary.append(f"情感分析置信度: {self.sentiment_df['sentiment_confidence'].mean():.3f}")
            summary.append("")

        if self.relationship_graph is not None:
            summary.append(f"关系网络: {self.relationship_graph.number_of_nodes()} 节点, "
                           f"{self.relationship_graph.number_of_edges()} 边")
            summary.append(f"网络密度: {nx.density(self.relationship_graph):.3f}")
            summary.append("")

        if self.topics is not None:
            summary.append(f"发现主题: {len(self.topics)} 个")
            for topic in self.topics[:3]:  # 只显示前3个主题
                summary.append(f"  • {topic['topic_type']}: {', '.join(topic['top_words'][:5])}")

        # 添加建议
        summary.append("")
        summary.append("改进建议:")
        if len(self.df) < 5:
            summary.append("1. 增加分析的科学家人数，建议至少5位")
        if self.sentiment_df is not None and self.sentiment_df['sentiment_confidence'].mean() < 0.6:
            summary.append("2. 考虑使用更专业的情感分析工具或扩充情感词典")
        if self.relationship_graph is not None and self.relationship_graph.number_of_edges() < 10:
            summary.append("3. 优化关系提取算法，提高关系识别准确率")

        with open(os.path.join(self.output_folder, "chinese_analysis_summary.txt"), 'w',
                  encoding='utf-8') as f:
            f.write('\n'.join(summary))

    def run_complete_analysis(self, enable_topic_modeling: bool = True) -> bool:
        """运行完整中文分析流程"""
        logger.info("=" * 60)
        logger.info("🚀 开始中文女科学家传记分析")
        logger.info("=" * 60)

        total_start = datetime.now()

        try:
            # 1. 加载和预处理
            if not self.load_and_preprocess_biographies():
                logger.error("数据加载失败")
                return False

            # 2. 情感分析
            if not self.analyze_sentiment_for_all():
                logger.warning("情感分析出现警告，继续执行...")

            # 3. 关系网络构建
            if not self.build_relationship_network():
                logger.warning("关系网络构建出现警告，继续执行...")

            # 4. 主题建模（可选）
            if enable_topic_modeling:
                self.perform_chinese_topic_modeling(num_topics=min(5, len(self.df)), method='lda')

            # 5. 可视化
            try:
                self.create_visualizations()
            except Exception as e:
                logger.warning(f"可视化失败: {e}")

            # 6. 导出结果
            self.export_results()

            # 总耗时
            total_time = (datetime.now() - total_start).total_seconds()

            print("\n" + "=" * 60)
            print("🎉 中文分析成功完成！")
            print(f"⏱️  总耗时: {total_time:.1f} 秒")

            if self.sentiment_df is not None:
                print(f"📊 平均情感分数: {self.sentiment_df['avg_sentiment'].mean():.3f}")

            if self.relationship_graph is not None:
                print(
                    f"🕸️  关系网络: {self.relationship_graph.number_of_nodes()}节点/{self.relationship_graph.number_of_edges()}边")

            print(f"📁 结果文件位于: {self.output_folder}")
            print("=" * 60)

            # 清理缓存
            self._cleanup()

            return True

        except Exception as e:
            logger.error(f"分析过程中出现严重错误: {e}", exc_info=True)
            print(f"\n❌ 分析失败: {e}")
            return False

    def _cleanup(self):
        """清理资源"""
        self._cache.clear()
        gc.collect()
        logger.info("资源清理完成")


# =============================================================================
# 使用示例
# =============================================================================

def main():
    """主函数 - 使用示例"""

    # 替换为你的Word文档文件夹路径或单个Word文件路径
    word_folder_path = "D:\\Project\\shuju\\.venv\\wenjian"  # 请修改为实际路径

    # 创建中文分析器实例
    analyzer = ChineseScientistBiographyAnalyzer(
        input_path=word_folder_path,
        output_folder="chinese_analysis_results"
    )

    # 运行完整中文分析
    success = analyzer.run_complete_analysis(
        enable_topic_modeling=True
    )

    if success:
        print("\n✅ 中文分析完成！")
        print("📋 主要功能:")
        print("  • 中文分词和文本预处理")
        print("  • 基于词典的中文情感分析")
        print("  • 中文关系提取和网络构建")
        print("  • 中文主题建模")
        print("  • 中文可视化图表")
    else:
        print("\n❌ 分析失败，请检查日志文件")


if __name__ == "__main__":
    # 安装必要依赖的提示
    print("=" * 60)
    print("中文女科学家传记分析系统")
    print("=" * 60)
    print("请确保已安装以下依赖:")
    print("  pip install jieba pandas numpy matplotlib seaborn networkx scikit-learn python-docx")
    print("=" * 60)

    main()